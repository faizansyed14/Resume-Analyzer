# app/services/parsing_service.py
import spacy
import spacy_transformers
import fitz
import re
from typing import Dict, Optional, Tuple, List
import concurrent.futures
import pytesseract
from PIL import Image
import io
import os
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor, as_completed
import multiprocessing
import queue
import time
from ..services.llm_service import LLMService
import docx
import pdfplumber
import phonenumbers
from email_validator import validate_email, EmailNotValidError
from nameparser import HumanName
from pathlib import Path
from typing import Union
from pyresparser import ResumeParser
from flask import Blueprint, request, jsonify, current_app
def safe_join(items: Union[str, List[str]]) -> str:
    if isinstance(items, list):
        return "\n".join(items)
    return str(items)

EMAIL_REGEX = re.compile(r"\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[A-Za-z]{2,}\b")
PHONE_REGEX = re.compile(r"(\+\d{1,3}[-.\s]?)?\d{3}[-.\s]?\d{3,4}[-.\s]?\d{4}")
URL_REGEX = re.compile(
    r"(https?://)?(www\.)?(linkedin\.com/in/|github\.com/|angel\.co/|bitbucket\.org/|behance\.net/|dribbble\.com/|stackoverflow\.com/users/|[a-z0-9.-]+\.(com|net|org|dev|me|tech|ai)/)[\w\d\-._~:/?#[\]@!$&'()*+,;=]*"
)
def get_llm_service():
    return current_app.llm_service_instance
class ParsingService:
    def __init__(self):
        try:
            self.nlp = spacy.load("en_core_web_trf")
        except OSError:
            print("en_core_web_trf not found, attempting to download...")
            spacy.cli.download("en_core_web_trf")
            self.nlp = spacy.load("en_core_web_trf")
        print("SpaCy transformer model loaded.")

        self.custom_headers = {
            "summary": ["summary", "professional summary", "profile", "objective", "career objective", "about me", "objective statement","skills summary", "SKILLS SUMMARY"],
            "skills": ["skills", "technical skills", "tools & technologies", "technologies used", "skills & abilities", "key competencies", "core strengths"],
            "experience": ["experience", "work experience", "professional experience", "employment history", "career history", "projects handled"],
            "education": ["education", "academic background", "qualifications", "education and training", "educational qualifications"],
            "certifications": ["certifications", "licenses", "certified", "courses completed", "certifications & achievements"],
            "languages": ["languages", "language proficiency", "known languages", "spoken languages"]
        }
    def _extract_section(self, text: str, headers: list[str]) -> str:
        """Extract a section from text based on headers"""
        if not text or not headers:
            return ""
            
        headers_combined = headers + [h.lower() for h in headers]
        headers_sorted = sorted(set(headers_combined), key=len, reverse=True)
        
        # Create pattern that matches any of the headers
        pattern = re.compile(
            r"(?i)(?:" + "|".join(re.escape(h) for h in headers_sorted) + r")\s*[:\-–]?\s*[\r\n]+"
        )
        
        # Find the section start
        match = pattern.search(text)
        if not match:
            return ""
            
        start = match.end()
        
        # Find the next section header (if any)
        next_header_pattern = re.compile(
            r"(?i)(?:" + "|".join(re.escape(h) for h in self._get_all_section_headers()) + r")\s*[:\-–]?\s*[\r\n]+"
        )
        
        next_match = next_header_pattern.search(text[start:])
        end = start + next_match.start() if next_match else len(text)
        
        return text[start:end].strip()

    def _get_all_section_headers(self) -> list:
        """Get all possible section headers from custom_headers"""
        all_headers = []
        for section in self.custom_headers.values():
            all_headers.extend(section)
        return list(set(all_headers))

    def parse_resume_pdf(self, pdf_path: str) -> Tuple[str, dict, Optional[Dict]]:
        """
        Parse resume from PDF or DOCX file
        1. Extract text via pdfplumber (fallback to OCR if empty)
        2. Delegate to parse_resume_text for entity extraction
        """
        text = ""
        ext = os.path.splitext(pdf_path)[1].lower()
        
        if ext == '.pdf':
            try:
                with pdfplumber.open(pdf_path) as pdf:
                    for page in pdf.pages:
                        pt = page.extract_text()
                        if pt:
                            text += pt + "\n"
            except Exception as e:
                print(f"Error extracting text from PDF: {e}")
                text = ""
        elif ext == '.docx':
            try:
                doc = docx.Document(pdf_path)
                text = "\n".join([para.text for para in doc.paragraphs if para.text])
            except Exception as e:
                print(f"Error reading DOCX file: {e}")
                text = ""
        
        text = text.strip()

        if not text:
            extractor = self.ContactInfoExtractor(self.nlp)
            text = extractor._try_ocr(pdf_path)

        # Now hand off to text-based parser
        return self.parse_resume_text(text)

    def parse_resume_text(self, raw_text: str) -> Tuple[str, dict, Optional[Dict]]:
        """
        Given a string of resume text, extract & refine entities with the LLM.
        Returns (original_text, extracted_entities_dict, llm_token_usage_dict).
        """
        refined_entities, token_usage = self._extract_and_refine_resume_entities(raw_text)
        return raw_text, refined_entities, token_usage

    def parse_job_pdf(self, pdf_path: str) -> Tuple[str, Dict, Optional[Dict]]:
        """Parse a job description PDF or DOCX file"""
        text = ""
        ext = os.path.splitext(pdf_path)[1].lower()
        
        if ext == '.pdf':
            text = self._load_pdf(pdf_path)
            if not text:
                print(f"Warning: No text extracted from {pdf_path}. Attempting OCR...")
                text = self._ocr_pdf(pdf_path)
        elif ext == '.docx':
            try:
                doc = docx.Document(pdf_path)
                text = "\n".join([para.text for para in doc.paragraphs if para.text])
            except Exception as e:
                print(f"Error reading DOCX file: {e}")
                text = ""
        
        if not text:
            return "", {"error": "Could not extract text from file"}, None

        refined_entities, llm_token_usage = self._extract_and_refine_job_entities(text)
        return text, refined_entities, llm_token_usage

    def parse_job_docx(self, docx_path: str) -> Tuple[str, Dict, Optional[Dict]]:
        """Parse a job description DOCX file"""
        try:
            doc = docx.Document(docx_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text])
            
            if not text:
                return "", {"error": "Could not extract text from DOCX file"}, None

            refined_entities, llm_token_usage = self._extract_and_refine_job_entities(text)
            return text, refined_entities, llm_token_usage
        except Exception as e:
            print(f"Error reading DOCX file: {e}")
            return "", {"error": str(e)}, None

    def parse_job_text(self, text: str) -> Tuple[str, Dict, Optional[Dict]]:
        """Parse job description text"""
        refined_entities, llm_token_usage = self._extract_and_refine_job_entities(text)
        return text, refined_entities, llm_token_usage

    def _load_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF using PyMuPDF (fitz)"""
        text = ""
        try:
            doc = fitz.open(pdf_path)
            for page in doc:
                text += page.get_text("text")
            doc.close()
            # Clean up the text
            text = re.sub(r'\s{2,}', ' ', text).strip()
            text = re.sub(r'(\n\s*){2,}', '\n\n', text)
            return text
        except Exception as e:
            print(f"Error reading PDF {pdf_path} with PyMuPDF: {e}")
            return ""

    def _ocr_pdf(self, pdf_path: str) -> str:
        """Fallback OCR for scanned PDFs"""
        text = ""
        try:
            doc = fitz.open(pdf_path)
            for page_num, page in enumerate(doc):
                pix = page.get_pixmap(dpi=300)
                img_bytes = pix.pil_tobytes(format="PNG")
                img = Image.open(io.BytesIO(img_bytes))
                # Preprocess image for better OCR
                img = img.convert('L')  # Grayscale
                img = img.point(lambda x: 0 if x < 160 else 255, '1')  # Threshold
                page_text = pytesseract.image_to_string(img)
                text += page_text + "\n\n"
            doc.close()
            # Clean up the text
            text = re.sub(r'\s{2,}', ' ', text).strip()
            text = re.sub(r'(\n\s*){2,}', '\n\n', text)
            return text
        except Exception as e:
            print(f"Exception during OCR of {pdf_path}: {e}")
            return ""

    def _parse_list(self, block: str) -> list[str]:
        """Parse a block of text into a list of items"""
        block_clean = block.replace('\xa0', ' ')
        lines = re.split(r"[\n;]|\s*(?<!\d),(?!\s*\d{4})", block_clean)
        seen = set()
        items = []
        for line in lines:
            s = line.strip()
            if not s:
                continue
            s = re.sub(r"^([•\-oO*•●–\d]+\s*)+", "", s).strip()
            s = re.sub(r"\s{2,}", " ", s)
            if s.lower() not in seen:
                seen.add(s.lower())
                items.append(s)
        return items

    class ContactInfoExtractor:
        def __init__(self, nlp):
            self.nlp = nlp
            # Enhanced email pattern
            self.email_pattern = re.compile(
                r"\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b",
                re.IGNORECASE
            )
            # More comprehensive phone pattern
            self.phone_pattern = re.compile(
                r"(?:\+?\d{1,3}[-.\s]?)?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}\b"
            )
            # Improved name pattern
            self.name_pattern = re.compile(
                r"^([A-Z][a-zA-Z'-]+(?:\s+[A-Z][a-zA-Z'-]+){1,3})(?:\s|$)",
                re.MULTILINE
            )
            
            # Expanded blacklists
            self.email_blacklist = {
                "sample@", "example@", "test@", "your@", "email@", 
                "info@", "contact@", "hello@", "name@"
            }
            self.phone_blacklist = {
                "1234567890", "123-456-7890", "0000000000",
                "5555555555", "1111111111", "0123456789"
            }

        def extract_from_text(self, text: str) -> Dict[str, Optional[str]]:
            """Enhanced extraction from text"""
            text = self._preprocess_text(text)
            
            results = {
                "name": self._extract_name(text),
                "email": self._extract_email(text),
                "phone": self._extract_phone(text),
                "linkedin": self._extract_linkedin(text)
            }
            
            return results

        def _preprocess_text(self, text: str) -> str:
            """Better text cleaning"""
            # Normalize whitespace and line breaks
            text = re.sub(r'\s+', ' ', text)
            # Remove common headers
            text = re.sub(r'(?i)\b(?:resume|cv|curriculum vitae)\b', '', text)
            return text.strip()

        def _extract_name(self, text: str) -> Optional[str]:
            """More reliable name extraction"""
            # First look for a clear name pattern at the start
            name_match = self.name_pattern.search(text)
            if name_match:
                candidate = name_match.group(1).strip()
                if self._is_valid_name(candidate):
                    return candidate.title()
            
            # Fallback to spaCy NER
            doc = self.nlp(text[:1000])  # Only check first part
            for ent in doc.ents:
                if ent.label_ == "PERSON" and self._is_valid_name(ent.text):
                    return ent.text.title()
            
            return None

        def _is_valid_name(self, name: str) -> bool:
            """Strict name validation"""
            if not name or len(name.split()) < 2:
                return False
            
            # Check for invalid patterns
            if any(char.isdigit() for char in name):
                return False
                
            # Check for common prefixes/suffixes
            first_word = name.split()[0].lower()
            if first_word in {'mr', 'mrs', 'ms', 'dr', 'prof'}:
                return False
                
            return True

        def _extract_email(self, text: str) -> Optional[str]:
            """Better email extraction"""
            emails = set()
            for match in self.email_pattern.finditer(text):
                email = match.group(0).lower()
                # Skip blacklisted patterns
                if not any(bl in email for bl in self.email_blacklist):
                    try:
                        valid = validate_email(email)
                        emails.add(valid.email)
                    except EmailNotValidError:
                        continue
            
            return next(iter(emails), None) if emails else None

        def _extract_phone(self, text: str) -> Optional[str]:
            """More reliable phone extraction"""
            phones = set()
            for match in self.phone_pattern.finditer(text):
                phone = match.group(0)
                clean_phone = re.sub(r'[^\d+]', '', phone)
                
                # Skip blacklisted numbers
                if clean_phone in self.phone_blacklist:
                    continue
                    
                try:
                    parsed = phonenumbers.parse(clean_phone, None)
                    if phonenumbers.is_valid_number(parsed):
                        formatted = phonenumbers.format_number(
                            parsed, 
                            phonenumbers.PhoneNumberFormat.E164
                        )
                        phones.add(formatted)
                except:
                    continue
            
            return next(iter(phones), None) if phones else None

        def _extract_linkedin(self, text: str) -> Optional[str]:
            """LinkedIn URL extraction"""
            linkedin_pattern = re.compile(
                r"(?:https?://)?(?:www\.)?linkedin\.com/(?:in|pub)/[a-zA-Z0-9\-]+/?",
                re.IGNORECASE
            )
            match = linkedin_pattern.search(text)
            return match.group(0) if match else None

    def _extract_resume_entities_initial(self, text: str) -> dict:
        """Improved initial extraction"""
        extractor = self.ContactInfoExtractor(self.nlp)
        contact_info = extractor.extract_from_text(text)
        
        # Get sections using the improved methods
        sections = {
            "summary_objective": self._extract_section(text, self.custom_headers["summary"]),
            "skills_raw": self._extract_section(text, self.custom_headers["skills"]),
            "experience_raw": self._extract_section(text, self.custom_headers["experience"]),
            "education": self._parse_list(self._extract_section(text, self.custom_headers["education"])),
            "certifications": self._parse_list(self._extract_section(text, self.custom_headers["certifications"]))
        }
        
        return {**contact_info, **sections}

    def _fallback_resume_extraction(self, text: str) -> dict:
        """Fallback extraction method if pyresparser fails"""
        print("\n=== Using Fallback Extraction ===")
        cleaned_text = text.replace('\xa0', ' ')
        doc = self.nlp(cleaned_text)
        entities = {
            "name": "",
            "email": "",
            "phone": "",
            "linkedin": "",
            "skills_raw": "",
            "experience_raw": "",
            "education": [],
            "summary_objective": "",
            "certifications": [],
            "projects": []
        }

        # Contact info
        m_email = EMAIL_REGEX.search(cleaned_text)
        if m_email: entities["email"] = m_email.group(0)

        m_phone = PHONE_REGEX.search(cleaned_text)
        if m_phone:
            raw = m_phone.group(0)
            cleaned_phone = re.sub(r"[^\d+]", "", raw)
            entities["phone"] = cleaned_phone

        m_linkedin = re.search(r"(linkedin\.com/in/[\w\d\-\._~]+)", cleaned_text, re.I)
        if m_linkedin: entities["linkedin"] = m_linkedin.group(0)

        # Name detection
        extractor = self.ContactInfoExtractor(self.nlp)
        entities["name"] = extractor._extract_name(cleaned_text)

        # Summary/Objective
        summary = self._extract_section(cleaned_text, self.custom_headers["summary"])
        entities["summary_objective"] = summary if summary else ""

        # Skills Section
        sk_raw = self._extract_section(cleaned_text, self.custom_headers["skills"])
        entities["skills_raw"] = sk_raw if sk_raw else ""

        # Experience Section
        ex_raw = self._extract_section(cleaned_text, self.custom_headers["experience"])
        entities["experience_raw"] = ex_raw if ex_raw else ""

        # Education Section
        ed_block = self._extract_section(cleaned_text, self.custom_headers["education"])
        if ed_block:
            entities["education"] = self._parse_list(ed_block)

        # Certifications Section
        cert_block = self._extract_section(cleaned_text, self.custom_headers["certifications"])
        if cert_block:
            entities["certifications"] = self._parse_list(cert_block)

        # Projects Section
        projects_block = self._extract_section(cleaned_text, ["projects", "personal projects", "key projects"])
        if projects_block:
            entities["projects"] = self._parse_list(projects_block)

        return entities

    def _extract_job_entities_initial(self, text: str) -> dict:
        """Extract job description entities using pyresparser"""
        print("\n=== Starting Job Description Extraction ===")
        
        try:
            # Create a temporary file for pyresparser
            temp_file = "temp_job.txt"
            with open(temp_file, "w", encoding="utf-8") as f:
                f.write(text)
            
            # Use pyresparser to extract entities
            parser = ResumeParser(temp_file)
            parsed_data = parser.get_extracted_data()
            
            # Map pyresparser output to our format
            entities = {
                "job_title": "",
                "company": "",
                "location": "",
                "skills_raw": "\n".join(parsed_data.get("skills", [])),
                "experience_raw": "\n".join([str(exp) for exp in parsed_data.get("experience", [])]),
                "education": parsed_data.get("degree", []),
                "languages": parsed_data.get("languages", [])
            }
            
            # Try to extract job title and company from first lines
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            if len(lines) > 0:
                entities["job_title"] = lines[0]
            if len(lines) > 1:
                entities["company"] = lines[1]
            
            return entities
            
        except Exception as e:
            print(f"Error in pyresparser for job: {e}")
            # Fallback to our own extraction
            return self._fallback_job_extraction(text)
        finally:
            try:
                os.remove(temp_file)
            except:
                pass

    def _fallback_job_extraction(self, text: str) -> dict:
        """Fallback job extraction method if pyresparser fails"""
        print("\n=== Using Fallback Job Extraction ===")
        cleaned_text = text.replace('\xa0', ' ')
        doc = self.nlp(text)
        entities = {
            "job_title": "",
            "company": "",
            "location": "",
            "skills_raw": "",
            "experience_raw": "",
            "education": [],
            "languages": []
        }

        # IMPROVED Job Title Extraction:
        role_keywords = [
            "engineer", "developer", "analyst", "architect", "consultant", "administrator",
            "manager", "specialist", "lead", "security", "network", "cloud", "data"
        ]

        # Search in the first 30 lines or 1000 characters
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        title_candidates = lines[:30] + re.split(r'\.|\n', text[:1000])

        for line in title_candidates:
            if any(kw in line.lower() for kw in role_keywords):
                if 2 <= len(line.split()) <= 8:
                    entities["job_title"] = line.strip()
                    break

        # Company and Location
        for ent_ in doc.ents:
            if ent_.start_char > 500:
                break
            if ent_.label_ in ("GPE", "LOC") and not entities["location"]:
                entities["location"] = ent_.text.strip()

        # Skills Section
        sk_raw = self._extract_section(text, ["skills", "key skills", "required skills"])
        if sk_raw:
            entities["skills_raw"] = sk_raw
        else:
            skill_matches = [sent.text.strip() for sent in doc.sents if re.search(r"skills?|technolog|tools", sent.text, re.I)]
            entities["skills_raw"] = " ".join(skill_matches)

        # Experience Section
        ex_raw = self._extract_section(text, ["experience", "responsibilities", "requirements"])
        if ex_raw:
            entities["experience_raw"] = ex_raw
        else:
            ex_sents = [sent.text.strip() for sent in doc.sents if re.search(r"\d+\+?\s+years|experience", sent.text, re.I)]
            entities["experience_raw"] = " ".join(ex_sents)

        # Education
        ed_block = self._extract_section(text, ["education", "qualifications"])
        if ed_block:
            for item in self._parse_list(ed_block):
                if re.search(r"\b(university|degree|bachelor|master|phd|diploma)\b", item, re.I):
                    entities["education"].append(item)
        entities["education"] = list(dict.fromkeys(entities["education"]))

        # Languages
        lang_block = self._extract_section(cleaned_text, self.custom_headers.get("languages", []))
        if lang_block:
            entities["languages"] = self._parse_list(lang_block)

        return entities

    def _extract_and_refine_resume_entities(self, text: str) -> Tuple[Dict, Optional[Dict]]:
        print("\n=== Starting Resume Entity Extraction ===")
        initial = self._extract_resume_entities_initial(text)
        
        final = {
            "name": initial.get("name", ""),
            "email": initial.get("email", ""),
            "phone": initial.get("phone", ""),
            "linkedin": initial.get("linkedin", ""),
            "summary_objective": initial.get("summary_objective", ""),
            "skills": self._parse_list(initial.get("skills_raw", "")),
            "experience": self._parse_list(initial.get("experience_raw", "")),
            "education": initial.get("education", []),
            "certifications": initial.get("certifications", []),
            "projects": initial.get("projects", [])
        }
        
        return final, None  # Return None for token usage since we're not using LLM here

    def _extract_and_refine_job_entities(self, text: str) -> Tuple[Dict, Optional[Dict]]:
        print("\n=== Starting Job Description Extraction ===")
        initial = self._extract_job_entities_initial(text)
        
        if not initial:  # Handle case where extraction fails
            return {}, None
            
        final = {
            "job_title": initial.get("job_title", ""),
            "company": initial.get("company", ""),
            "location": initial.get("location", ""),
            "skills": self._parse_list(initial.get("skills_raw", "")),
            "experience": self._parse_list(initial.get("experience_raw", "")),
            "education": initial.get("education", []),
            "summary_objective": initial.get("summary_objective", ""),
            "languages": initial.get("languages", [])
        }
        
        return final, None