# app/services/resume_parser.py
import os
import re
import io
import hashlib
import asyncio
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Dict, List, Tuple, Optional, Any
from multiprocessing import cpu_count
import time
from functools import lru_cache

# ---- File & OCR libs ----
import fitz  # PyMuPDF
import pdfplumber
import docx
import pytesseract
from PIL import Image

# ---- NLP libs ----
import spacy
from spacy.matcher import PhraseMatcher, Matcher

# ---- Validation libs ----
import phonenumbers
from email_validator import validate_email, EmailNotValidError

# ---- Embeddings / NER ----
from sentence_transformers import SentenceTransformer
from transformers import pipeline, AutoTokenizer, AutoModelForTokenClassification

# ---- Optional helpers (auto-fallback if missing) ----
try:
    from presidio_analyzer import AnalyzerEngine
    PRESIDIO_AVAILABLE = True
except Exception:
    PRESIDIO_AVAILABLE = False

try:
    from rapidfuzz import fuzz, process as fuzzprocess
    RAPIDFUZZ_AVAILABLE = True
except Exception:
    RAPIDFUZZ_AVAILABLE = False

try:
    import pdfminer.high_level
    from pdfminer.layout import LAParams
    PDFMINER_AVAILABLE = True
except Exception:
    PDFMINER_AVAILABLE = False

# Date parsing (robust ISO normalization)
try:
    from dateutil import parser as dateparse
    DATEUTIL_AVAILABLE = True
except Exception:
    DATEUTIL_AVAILABLE = False

# ---- Skill library (use your centralized database) ----
from .skill_library import SkillLibrary


class EnhancedResumeParser:
    """
    High-performance resume parser optimized for batch processing 100+ resumes.

    Fixes & improvements in this version:
      - Robust email/phone/LinkedIn extraction (handles obfuscations and noise)
      - Experience extraction does NOT rely only on "Experience" section
      - Education extraction with whole-document scan + GPA, dates, institution heuristics
      - Language extraction from anywhere (uses spoken_languages category)
      - Safer spaCy pipeline disabling (no E001)
      - Processing time is actual elapsed seconds
      - .txt file support
    """

    # ---------- Performance Config ----------
    OCR_DPI = 300
    OCR_PSM = 6
    OCR_OEM = 3
    MAX_WORKERS = min(8, cpu_count() or 2)
    BATCH_SIZE = 10
    MAX_NAME_SCAN_CHARS = 2000
    CACHE_SIZE = 1000

    # ---------- Processing Modes ----------
    SINGLE_THREADED = "single"
    MULTI_THREADED = "threaded"
    ASYNC_BATCH = "async"

    def __init__(self, processing_mode: str = MULTI_THREADED):
        self.processing_mode = processing_mode
        self._initialize_models()
        self._initialize_skill_library()
        self._initialize_patterns()
        self._cache: Dict[str, Dict] = {}

        print(f"Enhanced Resume Parser initialized in {processing_mode} mode")

    # =========================
    # Initialization
    # =========================
    def _initialize_models(self):
        """Initialize all ML models with error handling"""
        # spaCy transformer (prefer trf; fallback to sm)
        try:
            self.nlp = spacy.load("en_core_web_trf")
            to_disable = [p for p in ("lemmatizer", "textcat") if p in self.nlp.pipe_names]
            if to_disable:
                self.nlp.disable_pipes(*to_disable)
            print("SpaCy model (trf) loaded. Pipeline:", self.nlp.pipe_names)
        except Exception:
            try:
                self.nlp = spacy.load("en_core_web_sm")
                to_disable = [p for p in ("lemmatizer", "textcat") if p in self.nlp.pipe_names]
                if to_disable:
                    self.nlp.disable_pipes(*to_disable)
                print("SpaCy model (sm) loaded. Pipeline:", self.nlp.pipe_names)
            except Exception:
                import spacy.cli as spacy_cli
                spacy_cli.download("en_core_web_sm")
                self.nlp = spacy.load("en_core_web_sm")
                to_disable = [p for p in ("lemmatizer", "textcat") if p in self.nlp.pipe_names]
                if to_disable:
                    self.nlp.disable_pipes(*to_disable)
                print("SpaCy model (sm) downloaded & loaded. Pipeline:", self.nlp.pipe_names)

        # SentenceTransformer (optional for semantic matching)
        try:
            self.sentence_model = SentenceTransformer("all-mpnet-base-v2")
            print("SentenceTransformer model loaded successfully")
        except Exception as e:
            print(f"SentenceTransformer failed to load: {e}")
            self.sentence_model = None

        # HuggingFace NER
        try:
            self.ner_tokenizer = AutoTokenizer.from_pretrained(
                "dbmdz/bert-large-cased-finetuned-conll03-english"
            )
            self.ner_model = AutoModelForTokenClassification.from_pretrained(
                "dbmdz/bert-large-cased-finetuned-conll03-english"
            )
            self.ner_pipeline = pipeline(
                "ner",
                model=self.ner_model,
                tokenizer=self.ner_tokenizer,
                aggregation_strategy="simple",
                device=-1,  # CPU
            )
            print("HuggingFace NER pipeline loaded")
        except Exception as e:
            print(f"HF NER load failed: {e}")
            self.ner_pipeline = None

        # Presidio (optional)
        self.presidio = None
        if PRESIDIO_AVAILABLE:
            try:
                self.presidio = AnalyzerEngine()
                print("Presidio AnalyzerEngine loaded")
            except Exception as e:
                print(f"Presidio init failed: {e}")

    def _initialize_skill_library(self):
        """Initialize skill matching with optimizations"""
        self.skill_library = SkillLibrary()
        self.all_skills = self.skill_library.get_all_skills()

        # Case-insensitive mapping for fast lookup
        self.skills_lower_map = {s.lower(): s for s in self.all_skills}

        # Build optimized PhraseMatcher
        self.skill_matcher = PhraseMatcher(self.nlp.vocab, attr="LOWER")
        skill_patterns = [self.nlp.make_doc(skill) for skill in self.all_skills]
        self.skill_matcher.add("SKILLS", skill_patterns)

        # Context matcher
        self.skill_context_matcher = Matcher(self.nlp.vocab)
        self._build_skill_context_patterns()

        print(f"Skill library initialized with {len(self.all_skills)} skills")

    def _build_skill_context_patterns(self):
        tech_patterns = [
            [{"LOWER": {"IN": ["stack", "technologies", "tools", "languages"]}}, {"TEXT": ":"}],
            [{"LOWER": "proficient"}, {"LOWER": "in"}],
            [{"LOWER": {"IN": ["experience", "skilled", "expertise"]}}, {"LOWER": "with"}],
        ]
        for i, pattern in enumerate(tech_patterns):
            self.skill_context_matcher.add(f"TECH_CONTEXT_{i}", [pattern])

    def _initialize_patterns(self):
        """Initialize all regex patterns for performance"""
        self.headers = {
            "summary": [
                "summary", "professional summary", "profile", "objective", "career objective",
                "about me", "professional profile", "introduction", "overview",
                "summary of qualifications", "executive summary", "career profile",
                "professional overview", "key qualifications", "background",
            ],
            "skills": [
                "skills", "technical skills", "tools & technologies", "technologies used",
                "skills & abilities", "technical competencies", "core competencies",
                "programming languages", "stack", "expertise", "proficiencies",
                "technical expertise", "key skills", "relevant skills", "competencies",
            ],
            "experience": [
                "experience", "work experience", "professional experience", "employment history",
                "career history", "projects handled", "work history", "relevant experience",
                "professional background", "employment", "career summary", "work summary",
            ],
            "education": [
                "education", "academic background", "qualifications", "education and training",
                "academic history", "degrees", "academic experience", "educational background",
                "academic qualifications", "schooling", "university", "college",
            ],
            "certifications": [
                "certifications", "licenses", "awards", "courses", "certificates",
                "professional certifications", "training", "credentials", "achievements",
            ],
            "projects": [
                "projects", "key projects", "personal projects", "notable projects",
                "project experience", "relevant projects", "academic projects",
            ],
            "languages": [
                "languages", "language proficiency", "spoken languages", "language skills",
            ],
        }

        self._compile_patterns()

    def _compile_patterns(self):
        self._re_email = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
        # Broader phone detection is handled by phonenumbers, but keep a permissive seed regex for speed:
        self._re_phone_seed = re.compile(r"(\+?\d[\d\-\.\s\(\)]{6,}\d)")

        self._re_linkedin = re.compile(
            r'(?:https?://)?(?:[a-z]{0,3}\.)?linkedin\.com/(?:in|pub|profile|u/\w+/profile)/[^\s,;]+/?',
            flags=re.I
        )
        self._re_github = re.compile(r"(?:https?://)?(?:www\.)?github\.com/[\w\-_]+/?", flags=re.I)

        # Text cleaning patterns
        self._re_spaces = re.compile(r"[ \t]+")
        self._re_newlines = re.compile(r"\n{3,}")
        self._re_cr = re.compile(r"\r")
        self._re_bullets = re.compile(r"(?:\n\s*\n|•|▪|▫|◦|‣|⁃|-{2,})")

        # Date ranges like "Jan 2019 - Present", "2017–2019", "03/2018 - 11/2020"
        self._re_date_range = re.compile(
            r"(?P<start>(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{4}|\d{1,2}/\d{4}|\d{4})"
            r"\s*[-–—to]\s*"
            r"(?P<end>(?:Present|Current|Now)|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{4}|\d{1,2}/\d{4}|\d{4})",
            flags=re.I
        )

        # Job title & company header patterns
        self._re_job_patterns = [
            re.compile(
                r"^(?P<title>[^,\n]{5,80})\s*(?:at|@)\s*(?P<company>[^,\n]{2,80})"
                r"(?:,\s*(?P<location>[^,\n]{2,80}))?",
                flags=re.I | re.M
            ),
            re.compile(
                r"^(?P<company>[A-Z][^,\n]{2,80})\s*[-–—]\s*(?P<title>[^,\n]{3,80})"
                r"(?:,\s*(?P<location>[^,\n]{2,80}))?",
                flags=re.I | re.M
            ),
            re.compile(
                r"^(?P<title>[^,\n]{3,80}),\s*(?P<company>[A-Z][^,\n]{2,80})"
                r"(?:,\s*(?P<location>[^,\n]{2,80}))?",
                flags=re.I | re.M
            ),
        ]

        self._re_bullet_line = re.compile(r"^\s*(?:[-*•▪▫◦‣⁃]\s+|\d+\.\s+|→\s+)")

        # Education helpers
        self._re_degree_keywords = re.compile(
            r"(Bachelor|Master|B\.?Sc|M\.?Sc|B\.?E|M\.?E|B\.?Tech|M\.?Tech|BS|MS|Ph\.?D\.?|MBA|MFA|LLB|MD|JD|Associate|Diploma|Certificate)",
            re.I
        )

    # =========================
    # Main Processing API
    # =========================
    async def process_resumes_batch(self, file_paths: List[str]) -> List[Dict]:
        if self.processing_mode == self.ASYNC_BATCH:
            return await self._process_async_batch(file_paths)
        elif self.processing_mode == self.MULTI_THREADED:
            return await self._process_threaded_batch(file_paths)
        else:
            return await self._process_sequential_batch(file_paths)

    async def _process_async_batch(self, file_paths: List[str]) -> List[Dict]:
        results = []
        semaphore = asyncio.Semaphore(self.MAX_WORKERS)

        async def process_single(file_path: str) -> Dict:
            async with semaphore:
                try:
                    return await self._process_single_async(file_path)
                except Exception as e:
                    return {"file_path": file_path, "error": str(e), "entities": {}}

        for i in range(0, len(file_paths), self.BATCH_SIZE):
            batch = file_paths[i:i + self.BATCH_SIZE]
            batch_results = await asyncio.gather(*[process_single(fp) for fp in batch])
            results.extend(batch_results)
            if len(self._cache) > self.CACHE_SIZE:
                self._cache.clear()

        return results

    async def _process_threaded_batch(self, file_paths: List[str]) -> List[Dict]:
        loop = asyncio.get_event_loop()
        results: List[Dict] = []
        with ThreadPoolExecutor(max_workers=self.MAX_WORKERS) as executor:
            tasks = [loop.run_in_executor(executor, self._process_single_sync, fp) for fp in file_paths]
            for completed in asyncio.as_completed(tasks):
                try:
                    result = await completed
                except Exception as e:
                    result = {"file_path": "unknown", "error": str(e), "entities": {}}
                results.append(result)
        return results

    async def _process_sequential_batch(self, file_paths: List[str]) -> List[Dict]:
        results = []
        for file_path in file_paths:
            try:
                result = await self._process_single_async(file_path)
            except Exception as e:
                result = {"file_path": file_path, "error": str(e), "entities": {}}
            results.append(result)
        return results

    async def _process_single_async(self, file_path: str) -> Dict:
        cache_key = self._get_cache_key(file_path)
        if cache_key in self._cache:
            return self._cache[cache_key]

        t0 = time.perf_counter()
        raw_text = await self._extract_text_async(file_path)
        entities = await self._extract_entities_async(raw_text)
        elapsed = time.perf_counter() - t0

        result = {
            "file_path": file_path,
            "raw_text": raw_text,
            "entities": entities,
            "processing_time": round(elapsed, 3),
        }
        self._cache[cache_key] = result
        return result

    def _process_single_sync(self, file_path: str) -> Dict:
        try:
            cache_key = self._get_cache_key(file_path)
            if cache_key in self._cache:
                return self._cache[cache_key]

            t0 = time.perf_counter()
            raw_text = self._extract_text_sync(file_path)
            entities = self._extract_entities_sync(raw_text)
            elapsed = time.perf_counter() - t0

            result = {
                "file_path": file_path,
                "raw_text": raw_text,
                "entities": entities,
                "processing_time": round(elapsed, 3),
            }
            self._cache[cache_key] = result
            return result
        except Exception as e:
            return {"file_path": file_path, "error": str(e), "entities": {}}

    def _get_cache_key(self, file_path: str) -> str:
        try:
            stat = os.stat(file_path)
            content = f"{file_path}_{stat.st_mtime}_{stat.st_size}"
            return hashlib.md5(content.encode()).hexdigest()
        except Exception:
            return hashlib.md5(file_path.encode()).hexdigest()

    # =========================
    # Legacy API Compatibility
    # =========================
    def parse_resume_pdf(self, pdf_path: str) -> Tuple[str, dict, Optional[Dict]]:
        result = self._process_single_sync(pdf_path)
        return result.get("raw_text", ""), result.get("entities", {}), None

    def parse_resume_text(self, raw_text: str) -> Tuple[str, dict, Optional[Dict]]:
        cleaned = self._clean_text(raw_text or "")
        entities = self._extract_entities_sync(cleaned)
        return raw_text, entities, None

    # =========================
    # Text Extraction
    # =========================
    async def _extract_text_async(self, file_path: str) -> str:
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self._extract_text_sync, file_path)

    def _extract_text_sync(self, file_path: str) -> str:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()

        try:
            if ext == ".pdf":
                return self._extract_pdf_text_enhanced(file_path)
            elif ext == ".docx":
                return self._extract_docx_text_enhanced(file_path)
            elif ext == ".txt":
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    return self._clean_text(f.read())
            else:
                raise ValueError(f"Unsupported file type: {ext}")
        except Exception as e:
            print(f"Primary extraction failed for {file_path}: {e}")
            if ext == ".pdf":
                try:
                    return self._ocr_pdf_enhanced(file_path)
                except Exception as ocr_e:
                    print(f"OCR fallback failed: {ocr_e}")
            return ""

    def _extract_pdf_text_enhanced(self, pdf_path: str) -> str:
        best_text = ""
        best_score = 0
        for extractor in (self._extract_with_pdfplumber, self._extract_with_pymupdf, self._extract_with_pdfminer):
            try:
                text = extractor(pdf_path)
                score = self._calculate_text_quality(text)
                if score > best_score:
                    best_text, best_score = text, score
                if score > 0.8:
                    break
            except Exception as e:
                print(f"Extractor {extractor.__name__} failed: {e}")
        return best_text if best_score > 0.1 else self._ocr_pdf_enhanced(pdf_path)

    def _extract_with_pdfplumber(self, pdf_path: str) -> str:
        with pdfplumber.open(pdf_path) as pdf:
            parts = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    parts.append(page_text)
            return "\n".join(parts)

    def _extract_with_pymupdf(self, pdf_path: str) -> str:
        doc = fitz.open(pdf_path)
        parts = [page.get_text() for page in doc]
        doc.close()
        return "\n".join(parts)

    def _extract_with_pdfminer(self, pdf_path: str) -> str:
        if not PDFMINER_AVAILABLE:
            raise ImportError("pdfminer not available")
        laparams = LAParams(boxes_flow=0.5, word_margin=0.1, char_margin=2.0, line_margin=0.5)
        return pdfminer.high_level.extract_text(pdf_path, laparams=laparams)

    def _calculate_text_quality(self, text: str) -> float:
        if not text or len(text) < 50:
            return 0.0
        alpha_ratio = sum(c.isalpha() for c in text) / len(text)
        if alpha_ratio < 0.3:
            return 0.1
        words = re.findall(r'\b[a-zA-Z]{2,}\b', text)
        if len(words) < 20:
            return 0.2
        resume_keywords = ['experience', 'education', 'skills', 'work', 'university', 'company']
        keyword_count = sum(1 for kw in resume_keywords if kw.lower() in text.lower())
        base = min(alpha_ratio * 2, 1.0)
        bonus = min(keyword_count * 0.1, 0.3)
        return min(base + bonus, 1.0)

    def _extract_docx_text_enhanced(self, docx_path: str) -> str:
        try:
            document = docx.Document(docx_path)
            parts = []
            for p in document.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for table in document.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        parts.append(" | ".join(row_text))
            return self._clean_text("\n".join(parts))
        except Exception as e:
            print(f"DOCX extraction failed: {e}")
            return ""

    def _ocr_pdf_enhanced(self, pdf_path: str) -> str:
        try:
            doc = fitz.open(pdf_path)
            text_parts = []
            with ThreadPoolExecutor(max_workers=min(4, self.MAX_WORKERS)) as executor:
                futures = [executor.submit(self._ocr_single_page, doc[p]) for p in range(len(doc))]
                for f in as_completed(futures):
                    try:
                        t = f.result()
                        if t:
                            text_parts.append(t)
                    except Exception as e:
                        print(f"OCR page failed: {e}")
            doc.close()
            return "\n".join(text_parts)
        except Exception as e:
            print(f"OCR extraction failed: {e}")
            return ""

    def _ocr_single_page(self, page) -> str:
        try:
            pix = page.get_pixmap(dpi=self.OCR_DPI)
            img = Image.open(io.BytesIO(pix.tobytes()))
            img = img.convert("L")
            config = f'--oem {self.OCR_OEM} --psm {self.OCR_PSM} -l eng'
            return pytesseract.image_to_string(img, config=config).strip()
        except Exception as e:
            print(f"Single page OCR failed: {e}")
            return ""

    # =========================
    # Entities
    # =========================
    async def _extract_entities_async(self, text: str) -> Dict:
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self._extract_entities_sync, text)

    def _extract_entities_sync(self, text: str) -> Dict:
        if not text or len(text.strip()) < 10:
            return self._empty_entities()

        cleaned_text = self._clean_text(text)

        contact = self._extract_contact_info_enhanced(cleaned_text)
        summary = self._extract_summary_enhanced(cleaned_text)
        skills = self._extract_skills_enhanced(cleaned_text)

        # Experience: try section, then whole-doc fallback
        experience = self._extract_experience_enhanced(cleaned_text)
        if not experience:
            experience = self._extract_experience_fallback_scandoc(cleaned_text)

        # Education: try section, then whole-doc fallback
        education = self._extract_education_enhanced(cleaned_text)
        if not education:
            education = self._extract_education_fallback_scandoc(cleaned_text)

        certifications = self._extract_certifications_enhanced(cleaned_text)

        # Languages: section first, then global scan using library
        languages = self._extract_languages_enhanced(cleaned_text)
        if not languages:
            languages = self._extract_languages_from_anywhere(cleaned_text)

        return {
            "name": contact.get("name"),
            "email": contact.get("email"),
            "phone": contact.get("phone"),
            "linkedin": contact.get("linkedin"),
            "github": contact.get("github"),
            "portfolio": contact.get("portfolio"),
            "summary_objective": summary,
            "skills": skills,
            "experience": experience,
            "education": education,
            "certifications": certifications,
            "projects": self._extract_projects_enhanced(cleaned_text),
            "languages": languages,
        }

    def _empty_entities(self) -> Dict:
        return {
            "name": None,
            "email": None,
            "phone": None,
            "linkedin": None,
            "github": None,
            "portfolio": None,
            "summary_objective": "",
            "skills": [],
            "experience": [],
            "education": [],
            "certifications": [],
            "projects": [],
            "languages": [],
        }

    # =========================
    # Contact Info (improved)
    # =========================
    def _extract_contact_info_enhanced(self, text: str) -> Dict[str, Optional[str]]:
        scan_text = text[:self.MAX_NAME_SCAN_CHARS]

        result = {
            "name": None,
            "email": None,
            "phone": None,
            "linkedin": None,
            "github": None,
            "portfolio": None,
        }

        # Name (your existing approach kept, but make resilient)
        name_candidates = self._extract_name_candidates(scan_text)
        result["name"] = self._select_best_name(name_candidates)

        # Email (handles obfuscation like "name [at] domain [dot] com")
        result["email"] = self._extract_email_enhanced(text)

        # Phone (uses phonenumbers matcher across whole doc)
        result["phone"] = self._extract_phone_enhanced(text)

        # LinkedIn / GitHub / Portfolio
        result["linkedin"] = self._extract_url_enhanced(text, self._re_linkedin)
        result["github"] = self._extract_url_enhanced(text, self._re_github)
        result["portfolio"] = self._extract_portfolio_enhanced(text)

        return result

    def _extract_name_candidates(self, text: str) -> List[Tuple[str, float]]:
        candidates: List[Tuple[str, float]] = []

        # HF NER
        if self.ner_pipeline:
            try:
                for entity in self.ner_pipeline(text):
                    if entity.get("entity_group") == "PER":
                        name = entity.get("word", "").strip()
                        if self._is_valid_name(name):
                            candidates.append((name, float(entity.get("score", 0)) * 0.9))
            except Exception:
                pass

        # spaCy NER
        try:
            doc = self.nlp(text)
            for ent in doc.ents:
                if ent.label_ == "PERSON":
                    name = ent.text.strip()
                    if self._is_valid_name(name):
                        position_score = 1.0 - (ent.start_char / max(len(text), 1))
                        candidates.append((name, position_score * 0.8))
        except Exception:
            pass

        # Presidio
        if self.presidio:
            try:
                presidio_results = self.presidio.analyze(text=text, language="en")
                for r in presidio_results:
                    if r.entity_type == "PERSON":
                        name = text[r.start:r.end].strip()
                        if self._is_valid_name(name):
                            candidates.append((name, float(r.score) * 0.7))
            except Exception:
                pass

        # Heuristic: first lines that look like a name
        for i, line in enumerate(text.split('\n')[:10]):
            line = line.strip()
            if self._is_potential_name_line(line):
                position_score = 1.0 - (i / 10)
                candidates.append((line, position_score * 0.6))

        return candidates

    def _is_valid_name(self, name: str) -> bool:
        if not name or len(name) < 2:
            return False
        name = re.sub(r'^(Mr\.?|Ms\.?|Mrs\.?|Dr\.?|Prof\.?)\s+', '', name, flags=re.I).strip()
        if any(ch.isdigit() for ch in name):
            return False
        if '@' in name:
            return False
        words = name.split()
        if not (2 <= len(words) <= 4):
            return False
        return all(w[0].isupper() and w.replace('-', '').isalpha() for w in words)

    def _is_potential_name_line(self, line: str) -> bool:
        if any(k in line.lower() for k in ['email', 'phone', 'address', 'linkedin', 'github']):
            return False
        if len(re.findall(r'[^a-zA-Z\s\-]', line)) > max(1, int(len(line) * 0.2)):
            return False
        return self._is_valid_name(line)

    def _select_best_name(self, candidates: List[Tuple[str, float]]) -> Optional[str]:
        if not candidates:
            return None
        candidates.sort(key=lambda x: x[1], reverse=True)
        best_name, best_score = candidates[0]
        if best_score <= 0.3:
            return None
        return self._clean_name(best_name)

    def _clean_name(self, name: str) -> str:
        name = re.sub(r'^(Mr\.?|Ms\.?|Mrs\.?|Dr\.?|Prof\.?)\s+', '', name, flags=re.I)
        words = [w.capitalize() for w in name.split() if all(c.isalpha() or c == '-' for c in w)]
        return ' '.join(words)

    # ---- Email / Phone / URLs ----
    def _deobfuscate_emails_text(self, text: str) -> str:
        # Replace common obfuscations: "name [at] domain [dot] com"
        patterns = [
            (r"\s*\[\s*at\s*\]\s*", "@"),
            (r"\s*\(\s*at\s*\)\s*", "@"),
            (r"\s+at\s+", "@"),
            (r"\s*\[\s*dot\s*\]\s*", "."),
            (r"\s*\(\s*dot\s*\)\s*", "."),
            (r"\s+dot\s+", "."),
        ]
        out = text
        for pat, rep in patterns:
            out = re.sub(pat, rep, out, flags=re.I)
        return out

    def _extract_email_enhanced(self, text: str) -> Optional[str]:
        # Try normal
        emails = self._re_email.findall(text)

        # Try deobfuscated
        if not emails:
            deob = self._deobfuscate_emails_text(text)
            emails = self._re_email.findall(deob)

        # Clean up mailto trailing punctuation
        cleaned = []
        for e in emails:
            e2 = e.strip().strip('.,;:()[]{}<>')
            try:
                valid = validate_email(e2)
                cleaned.append(valid.email)
            except EmailNotValidError:
                continue

        # Prefer non-example domains and the first in the document order
        prefer = [e for e in cleaned if not e.lower().endswith(("example.com", "email.com", "domain.com"))]
        return prefer[0] if prefer else (cleaned[0] if cleaned else None)

    def _extract_phone_enhanced(self, text: str) -> Optional[str]:
        # Use phonenumbers matcher across the whole document
        candidates = []
        for match in phonenumbers.PhoneNumberMatcher(text, None):  # region None = try all
            num = match.number
            if phonenumbers.is_valid_number(num):
                # Score longer national numbers higher
                nsn_len = len(phonenumbers.format_number(num, phonenumbers.PhoneNumberFormat.E164))
                candidates.append((num, nsn_len, match.start))
        if not candidates:
            # Seed with a permissive regex, then parse
            for m in self._re_phone_seed.finditer(text):
                s = m.group(0)
                for region in (None, 'US', 'IN', 'GB', 'CA'):
                    try:
                        num = phonenumbers.parse(s, region)
                        if phonenumbers.is_valid_number(num):
                            nsn_len = len(phonenumbers.format_number(num, phonenumbers.PhoneNumberFormat.E164))
                            candidates.append((num, nsn_len, m.start()))
                            break
                    except Exception:
                        continue

        if not candidates:
            return None

        # Sort by E164 length desc, then by earliest appearance
        candidates.sort(key=lambda t: (-t[1], t[2]))
        best = candidates[0][0]
        return phonenumbers.format_number(best, phonenumbers.PhoneNumberFormat.INTERNATIONAL)

    def _extract_url_enhanced(self, text: str, pattern: re.Pattern) -> Optional[str]:
        match = pattern.search(text)
        if not match:
            return None
        url = match.group(0)
        url = url.rstrip(').,;:')  # strip trailing punctuation
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        return url

    def _extract_portfolio_enhanced(self, text: str) -> Optional[str]:
        # Prefer a non-social personal domain if present
        portfolio_pattern = re.compile(
            r'(?:https?://)?(?:www\.)?'
            r'(?!(?:linkedin|github|facebook|twitter|instagram|youtube|medium)\.com)'
            r'([a-zA-Z0-9-]+\.(?:com|net|org|dev|me|io|co|portfolio))'
            r'(?:/[^\s,;]*)?',
            flags=re.I
        )
        match = portfolio_pattern.search(text)
        if match:
            url = match.group(0).rstrip(').,;:')
            if not url.startswith(('http://', 'https://')):
                url = 'https://' + url
            return url
        return None

    # =========================
    # Skills (unchanged core, validated with library)
    # =========================
    def _extract_skills_enhanced(self, text: str) -> List[str]:
        if not text:
            return []

        skills_from_section = self._extract_skills_from_section(text)
        skills_from_nlp = self._extract_skills_with_nlp(text)
        skills_from_context = self._extract_skills_from_context(text)

        all_skills = set(skills_from_section) | set(skills_from_nlp) | set(skills_from_context)
        validated = self._validate_and_filter_skills(list(all_skills), text)
        return sorted(validated, key=lambda x: x.lower())

    def _extract_skills_from_section(self, text: str) -> List[str]:
        skills_section = self._extract_section(text, self.headers["skills"])
        if not skills_section:
            return []
        skills = []
        items = re.split(r'[,;|•▪▫◦‣⁃\n\t]', skills_section)
        for item in items:
            item = item.strip(' -•▪▫◦‣⁃\t\n')
            if 2 <= len(item) <= 50:
                canonical = self._find_canonical_skill(item)
                if canonical:
                    skills.append(canonical)
        return skills

    def _extract_skills_with_nlp(self, text: str) -> List[str]:
        try:
            doc = self.nlp(text)
            matches = self.skill_matcher(doc)
            found = []
            for _, start, end in matches:
                span = doc[start:end]
                if self._is_valid_skill_context(doc, start, end):
                    canonical = self._find_canonical_skill(span.text.strip())
                    if canonical:
                        found.append(canonical)
            return list(set(found))
        except Exception as e:
            print(f"NLP skill extraction failed: {e}")
            return []

    def _extract_skills_from_context(self, text: str) -> List[str]:
        skills = []
        stack_patterns = [
            r'(?:technologies?|tools?|stack|languages?)[:\s]*([^.\n]{20,200})',
            r'(?:experience|skilled|proficient)\s+(?:in|with)[:\s]*([^.\n]{10,100})',
            r'(?:using|utilized|worked\s+with)[:\s]*([^.\n]{10,100})',
        ]
        for pat in stack_patterns:
            for m in re.finditer(pat, text, re.I):
                chunk = m.group(1)
                pieces = re.split(r'[,;|&\+/]', chunk)
                for s in pieces:
                    s = s.strip()
                    if 2 <= len(s) <= 30:
                        canonical = self._find_canonical_skill(s)
                        if canonical:
                            skills.append(canonical)
        return skills

    def _is_valid_skill_context(self, doc, start: int, end: int) -> bool:
        context = doc[max(0, start - 5): min(len(doc), end + 5)]
        neg = {'no', 'not', 'never', 'without', 'lack', 'avoid'}
        return not any(w in context.text.lower() for w in neg)

    def _find_canonical_skill(self, skill_text: str) -> Optional[str]:
        skill_text = skill_text.strip()
        if not skill_text:
            return None
        canonical = self.skills_lower_map.get(skill_text.lower())
        if canonical:
            return canonical
        normalized = self.skill_library.normalize_alias(skill_text)
        if normalized:
            return normalized
        if RAPIDFUZZ_AVAILABLE and len(skill_text) >= 3:
            best = fuzzprocess.extractOne(skill_text, self.all_skills, scorer=fuzz.WRatio, score_cutoff=85)
            if best:
                return best[0]
        return None

    def _validate_and_filter_skills(self, skills: List[str], full_text: str) -> List[str]:
        out = []
        for s in skills:
            if not s:
                continue
            if self.skill_library.is_ambiguous(s) and not self._has_strong_skill_context(s, full_text):
                continue
            if self._skill_mentioned_in_text(s, full_text):
                out.append(s)
        return list(set(out))

    def _has_strong_skill_context(self, skill: str, text: str) -> bool:
        pats = [
            rf'\b{re.escape(skill)}\s+(?:experience|skills?|proficiency)\b',
            rf'(?:expert|proficient|experienced)\s+in\s+{re.escape(skill)}',
            rf'{re.escape(skill)}\s+(?:certification|training|course)\b',
        ]
        return any(re.search(p, text, re.I) for p in pats)

    def _skill_mentioned_in_text(self, skill: str, text: str) -> bool:
        if re.search(rf'\b{re.escape(skill)}\b', text, re.I):
            return True
        variations = []
        if '.' in skill:
            variations.append(skill.replace('.', ''))
        if ' ' in skill:
            variations.append(skill.replace(' ', ''))
            variations.append(skill.replace(' ', '-'))
        for v in variations:
            if re.search(rf'\b{re.escape(v)}\b', text, re.I):
                return True
        return False

    # =========================
    # Experience (improved)
    # =========================
    def _extract_experience_enhanced(self, text: str) -> List[Dict]:
        section = self._extract_section(text, self.headers["experience"])
        if not section:
            return []
        return self._parse_experience_section(section)

    def _parse_experience_section(self, section: str) -> List[Dict]:
        blocks = self._split_experience_blocks(section)
        experiences = []
        for block in blocks:
            if len(block.strip()) < 20:
                continue
            exp = self._parse_experience_block(block)
            if exp and self._is_valid_experience(exp):
                experiences.append(exp)
        return experiences

    def _split_experience_blocks(self, section: str) -> List[str]:
        separators = [
            r'\n\s*\n',  # Blank line
            r'\n(?=[A-Z][^\n]{3,100}(?:\s+at\s+|\s*[-–—]\s*)[A-Z])',  # Title at Company
            r'\n(?=\d{1,2}/\d{4}\s*[-–—])',  # Date start
        ]
        blocks = [section]
        for sep in separators:
            new_blocks = []
            for b in blocks:
                parts = re.split(sep, b)
                new_blocks.extend([p.strip() for p in parts if p.strip()])
            blocks = new_blocks
        return blocks

    def _parse_experience_block(self, block: str) -> Dict:
        lines = [ln.strip() for ln in block.split('\n') if ln.strip()]
        if not lines:
            return {}

        header_text = ' '.join(lines[:2])
        header_info = self._parse_job_header(header_text)
        dates_info = self._extract_dates_from_block(block)
        bullets = self._extract_bullets_enhanced(block)
        sentences = self._extract_sentences_enhanced(block)

        # Use NER to fill missing company/location
        if not header_info.get('company') or not header_info.get('location'):
            ner = self._extract_ner_from_block(block)
            header_info['company'] = header_info.get('company') or ner.get('company', '')
            header_info['location'] = header_info.get('location') or ner.get('location', '')

        return {
            "title": header_info.get('title', ''),
            "company": header_info.get('company', ''),
            "location": header_info.get('location', ''),
            "start_date": dates_info.get('start_date'),
            "end_date": dates_info.get('end_date'),
            "is_current": dates_info.get('is_current', False),
            "duration": dates_info.get('duration', ''),
            "bullets": bullets,
            "sentences": sentences,
            "description": self._create_description(bullets, sentences),
            "raw_block": block.strip(),
        }

    def _parse_job_header(self, header_text: str) -> Dict:
        out = {"title": "", "company": "", "location": ""}
        for pat in self._re_job_patterns:
            m = pat.search(header_text.strip())
            if m:
                g = m.groupdict()
                out["title"] = (g.get("title") or "").strip(" ,-")
                out["company"] = (g.get("company") or "").strip(" ,-")
                out["location"] = (g.get("location") or "").strip(" ,-")
                break
        return out

    def _extract_experience_fallback_scandoc(self, text: str) -> List[Dict]:
        """
        Whole-document fallback:
        - Find all date ranges
        - For each, look backward for a plausible "title/company" header line
        - Pull bullets/sentences below until next blank line or next date header
        """
        experiences: List[Dict] = []

        # Pre-split lines and map char offsets
        lines = text.split("\n")
        offsets = []
        pos = 0
        for ln in lines:
            offsets.append((pos, pos + len(ln)))
            pos += len(ln) + 1  # include \n

        for m in self._re_date_range.finditer(text):
            start_idx = m.start()
            # Find which line this is on
            line_idx = 0
            for i, (s, e) in enumerate(offsets):
                if s <= start_idx <= e:
                    line_idx = i
                    break

            # Look up 2 lines for header candidate
            header_chunk = " ".join(lines[max(0, line_idx - 2): line_idx + 1])
            header_info = self._parse_job_header(header_chunk)

            # If header weak, try NER on those lines
            if not header_info.get("company") or not header_info.get("title"):
                header_text = " ".join(lines[max(0, line_idx - 3): line_idx + 1])
                ner_fill = self._extract_ner_from_block(header_text)
                header_info["company"] = header_info.get("company") or ner_fill.get("company", "")
                header_info["location"] = header_info.get("location") or ner_fill.get("location", "")

            # Capture block below until next strong separator
            block_lines = [lines[line_idx]]
            j = line_idx + 1
            while j < len(lines):
                if not lines[j].strip():
                    # stop at double blank line
                    if j + 1 < len(lines) and not lines[j + 1].strip():
                        break
                # stop if a new big header appears
                if re.match(r'^\s*[A-Z][A-Z\s&/]{5,40}\s*[:\-–]?\s*$', lines[j]):
                    break
                # stop if next date range line
                if self._re_date_range.search(lines[j]):
                    break
                block_lines.append(lines[j])
                j += 1

            block = "\n".join(block_lines)
            dates = self._extract_dates_from_block(block)
            bullets = self._extract_bullets_enhanced(block)
            sentences = self._extract_sentences_enhanced(block)

            exp = {
                "title": header_info.get('title', ''),
                "company": header_info.get('company', ''),
                "location": header_info.get('location', ''),
                "start_date": dates.get('start_date'),
                "end_date": dates.get('end_date'),
                "is_current": dates.get('is_current', False),
                "duration": dates.get('duration', ''),
                "bullets": bullets,
                "sentences": sentences,
                "description": self._create_description(bullets, sentences),
                "raw_block": block.strip(),
            }
            if self._is_valid_experience(exp):
                experiences.append(exp)

        # Deduplicate similar entries (by title+company+start)
        uniq = []
        seen = set()
        for e in experiences:
            key = (e["title"].lower(), e["company"].lower(), e.get("start_date"))
            if key in seen:
                continue
            seen.add(key)
            uniq.append(e)
        return uniq

    def _extract_dates_from_block(self, block: str) -> Dict:
        result = {"start_date": None, "end_date": None, "is_current": False, "duration": ""}
        m = self._re_date_range.search(block)
        if m:
            start_raw = m.group('start')
            end_raw = m.group('end')
            result["start_date"] = self._normalize_date(start_raw)
            if re.search(r'(?i)present|current|now', end_raw or ''):
                result["is_current"] = True
                result["end_date"] = None
            else:
                result["end_date"] = self._normalize_date(end_raw)
            if result["start_date"] and (result["end_date"] or result["is_current"]):
                result["duration"] = self._calculate_duration(result["start_date"], result["end_date"], result["is_current"])
        return result

    def _normalize_date(self, date_str: str) -> Optional[str]:
        if not date_str:
            return None
        if DATEUTIL_AVAILABLE:
            try:
                parsed = dateparse.parse(date_str, default=dateparse.parse("2000-01-01"))
                return f"{parsed.year:04d}-{parsed.month:02d}"
            except Exception:
                pass
        m = re.search(r'\b(19|20)\d{2}\b', date_str)
        if m:
            return f"{m.group(0)}-01"
        return None

    def _calculate_duration(self, start_date: str, end_date: Optional[str], is_current: bool) -> str:
        try:
            if not DATEUTIL_AVAILABLE:
                return ""
            start = dateparse.parse(f"{start_date}-01")
            if is_current:
                from datetime import datetime
                end = datetime.now()
            elif end_date:
                end = dateparse.parse(f"{end_date}-01")
            else:
                return ""
            months = (end.year - start.year) * 12 + end.month - start.month
            if months < 12:
                return f"{months} month{'s' if months != 1 else ''}"
            years = months // 12
            rem = months % 12
            s = f"{years} year{'s' if years != 1 else ''}"
            if rem:
                s += f" {rem} month{'s' if rem != 1 else ''}"
            return s
        except Exception:
            return ""

    def _extract_bullets_enhanced(self, block: str) -> List[str]:
        bullets = []
        for line in block.split('\n'):
            if self._re_bullet_line.match(line):
                clean = re.sub(self._re_bullet_line, '', line).strip()
                clean = self._clean_bullet_text(clean)
                if self._is_valid_bullet(clean):
                    bullets.append(clean)
        return bullets

    def _clean_bullet_text(self, text: str) -> str:
        text = re.sub(r'\s+', ' ', text).strip()
        if text.endswith(('.', ':', ';')):
            text = text[:-1].strip()
        return text

    def _is_valid_bullet(self, bullet: str) -> bool:
        if not bullet or len(bullet) < 10:
            return False
        if re.match(r'^\d{4}[-/]\d{2}', bullet):
            return False
        return len(bullet.split()) >= 3

    def _extract_sentences_enhanced(self, block: str) -> List[str]:
        try:
            clean = re.sub(self._re_bullet_line, '', block)
            clean = re.sub(r'\s+', ' ', clean).strip()
            doc = self.nlp(clean)
            out = []
            for sent in doc.sents:
                s = sent.text.strip()
                if self._is_valid_sentence(s):
                    out.append(s)
            return out
        except Exception as e:
            print(f"Sentence extraction failed: {e}")
            parts = re.split(r'[.!?]+', block)
            return [s.strip() for s in parts if s.strip() and len(s.strip()) > 10]

    def _is_valid_sentence(self, s: str) -> bool:
        if not s or len(s) < 10:
            return False
        if re.match(r'^\d{4}[-/]', s):
            return False
        words = s.split()
        return 3 <= len(words) <= 50

    def _create_description(self, bullets: List[str], sentences: List[str]) -> str:
        out = []
        for b in bullets[:5]:
            out.append(f"• {b}")
        for s in sentences[:3]:
            if not any(self._text_similarity(s, b) > 0.7 for b in bullets):
                out.append(s)
        return '\n'.join(out)

    def _text_similarity(self, t1: str, t2: str) -> float:
        if not t1 or not t2:
            return 0.0
        w1, w2 = set(t1.lower().split()), set(t2.lower().split())
        if not w1 or not w2:
            return 0.0
        inter = len(w1 & w2)
        union = len(w1 | w2)
        return inter / union if union else 0.0

    def _extract_ner_from_block(self, block: str) -> Dict:
        out = {"company": "", "location": ""}
        try:
            doc = self.nlp(block[:200])
            for ent in doc.ents:
                if ent.label_ == "ORG" and not out["company"]:
                    if self._looks_like_company(ent.text):
                        out["company"] = ent.text.strip()
                elif ent.label_ in ("GPE", "LOC") and not out["location"]:
                    out["location"] = ent.text.strip()
        except Exception as e:
            print(f"NER extraction failed: {e}")
        return out

    def _looks_like_company(self, text: str) -> bool:
        if not text or len(text) < 2:
            return False
        bad = {'experience', 'skills', 'education', 'summary', 'objective'}
        if text.lower() in bad:
            return False
        return text[0].isupper()

    def _is_valid_experience(self, e: Dict) -> bool:
        if not e.get('title') and not e.get('company'):
            return False
        if not e.get('bullets') and not e.get('sentences'):
            return False
        return True

    # =========================
    # Education (improved)
    # =========================
    def _extract_education_enhanced(self, text: str) -> List[Dict]:
        section = self._extract_section(text, self.headers["education"])
        if not section:
            return []
        return self._parse_education_section(section)

    def _parse_education_section(self, section: str) -> List[Dict]:
        blocks = self._split_education_blocks(section)
        out = []
        for b in blocks:
            if len(b.strip()) < 15:
                continue
            entry = self._parse_education_block(b)
            if entry and self._is_valid_education(entry):
                out.append(entry)
        return out

    def _split_education_blocks(self, section: str) -> List[str]:
        separators = [
            r'\n\s*\n',
            r'\n(?=(?:Bachelor|Master|PhD|B\.?[ASE]\.?|M\.?[ASE]\.?|Ph\.?D\.?|MBA|Diploma|Certificate))',
        ]
        blocks = [section]
        for sep in separators:
            new_blocks = []
            for b in blocks:
                parts = re.split(sep, b, flags=re.I)
                new_blocks.extend([p.strip() for p in parts if p.strip()])
            blocks = new_blocks
        return blocks

    def _parse_education_block(self, block: str) -> Dict:
        result = {
            "degree": None,
            "field": None,
            "institution": None,
            "location": None,
            "dates": [],
            "duration": "",
            "gpa": None,
            "details": block.strip(),
        }

        result["degree"] = self._extract_degree_enhanced(block)
        result["field"] = self._extract_field_of_study(block)
        result["institution"] = self._extract_institution_enhanced(block)
        dates = self._extract_dates_from_education(block)
        result["dates"] = dates
        result["duration"] = self._format_education_duration(dates)
        result["gpa"] = self._extract_gpa(block)
        result["location"] = self._extract_location_from_education(block)
        return result

    def _extract_education_fallback_scandoc(self, text: str) -> List[Dict]:
        """
        Whole-document scan for degree/institution lines even when no 'Education' header exists.
        """
        entries: List[Dict] = []
        lines = [ln for ln in text.split('\n') if ln.strip()]
        i = 0
        while i < len(lines):
            line = lines[i]
            if self._re_degree_keywords.search(line):
                # Grab a small window around the degree line
                chunk_lines = lines[max(0, i - 1): min(len(lines), i + 4)]
                chunk = "\n".join(chunk_lines)
                entry = self._parse_education_block(chunk)
                if self._is_valid_education(entry):
                    entries.append(entry)
                    i += 3
                    continue
            i += 1

        # Deduplicate by (degree, institution)
        uniq, seen = [], set()
        for e in entries:
            key = (str(e.get("degree")).lower(), str(e.get("institution")).lower())
            if key in seen:
                continue
            seen.add(key)
            uniq.append(e)
        return uniq

    def _extract_degree_enhanced(self, text: str) -> Optional[str]:
        patterns = [
            r'(Bachelor(?:\s+of\s+[A-Za-z &]+)?)',
            r'(Master(?:\s+of\s+[A-Za-z &]+)?)',
            r'(Ph\.?D\.?(?:\s+in\s+[A-Za-z &]+)?)',
            r'(B\.?\s?(?:Sc|E|Tech)\.?(?:\s+[A-Za-z &]+)?)',
            r'(M\.?\s?(?:Sc|E|Tech)\.?(?:\s+[A-Za-z &]+)?)',
            r'(MBA|MFA|LLB|MD|JD)',
            r'(Associate(?:\s+of\s+[A-Za-z &]+)?)',
            r'(Diploma(?:\s+in\s+[A-Za-z &]+)?)',
            r'(Certificate(?:\s+in\s+[A-Za-z &]+)?)',
        ]
        for p in patterns:
            m = re.search(p, text, re.I)
            if m:
                return m.group(1).strip()
        return None

    def _extract_field_of_study(self, text: str) -> Optional[str]:
        pats = [
            r'(?:Bachelor|Master|PhD|B\.?[ASE]\.?|M\.?[ASE]\.?|MBA)\s+(?:of\s+)?(?:Science|Arts)?\s*(?:in\s+)?([^,\n.]{3,60})',
            r'(?:Major|Concentration|Specialization)\s*:\s*([^,\n.]{3,60})',
            r'Field\s*:\s*([^,\n.]{3,60})',
        ]
        for p in pats:
            m = re.search(p, text, re.I)
            if m:
                field = m.group(1).strip()
                field = re.sub(r'^(in|of)\s+', '', field, flags=re.I)
                if len(field) > 2:
                    return field
        return None

    def _extract_institution_enhanced(self, text: str) -> Optional[str]:
        try:
            doc = self.nlp(text[:250])
            for ent in doc.ents:
                if ent.label_ == "ORG" and self._looks_like_institution(ent.text):
                    return ent.text.strip()
        except Exception:
            pass
        for line in text.split('\n'):
            line = line.strip()
            if self._looks_like_institution(line):
                return line
        return None

    def _looks_like_institution(self, text: str) -> bool:
        if not text or len(text) < 5:
            return False
        keywords = ['university', 'college', 'institute', 'school', 'academy',
                    'polytechnic', 'seminary', 'conservatory']
        t = text.lower()
        if any(k in t for k in keywords):
            return True
        words = [w for w in text.split() if w.isalpha()]
        return len(words) >= 2 and all(w[0].isupper() for w in words)

    def _extract_dates_from_education(self, text: str) -> List[str]:
        dates = [m.group(0) for m in self._re_date_range.finditer(text)]
        year_pat = re.compile(r'\b(19|20)\d{2}\b')
        for m in year_pat.finditer(text):
            y = m.group(0)
            if y not in [d for d in dates if y in d]:
                dates.append(y)
        month_year_pat = re.compile(r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+(19|20)\d{2}\b', re.I)
        for m in month_year_pat.finditer(text):
            s = m.group(0)
            if not any(s in d for d in dates):
                dates.append(s)
        return dates

    def _format_education_duration(self, dates: List[str]) -> str:
        if not dates:
            return ""
        if len(dates) == 1:
            return dates[0]
        sorted_dates = sorted(dates)
        return f"{sorted_dates[0]} - {sorted_dates[-1]}"

    def _extract_gpa(self, text: str) -> Optional[str]:
        pats = [
            r'GPA:?\s*([0-4]\.\d{1,2})',
            r'Grade:?\s*([0-4]\.\d{1,2})',
            r'CGPA:?\s*([0-4]\.\d{1,2})',
            r'(\d\.\d{1,2})\s*/\s*4\.0',
        ]
        for p in pats:
            m = re.search(p, text, re.I)
            if m:
                try:
                    gpa = float(m.group(1))
                    if 0 <= gpa <= 4.0:
                        return f"{gpa:.2f}"
                except Exception:
                    continue
        return None

    def _extract_location_from_education(self, text: str) -> Optional[str]:
        try:
            doc = self.nlp(text)
            for ent in doc.ents:
                if ent.label_ in ("GPE", "LOC"):
                    return ent.text.strip()
        except Exception:
            pass
        return None

    def _is_valid_education(self, e: Dict) -> bool:
        if not e.get('degree') and not e.get('institution'):
            return False
        return len((e.get('details') or '').strip()) > 10

    # =========================
    # Summary / Certifications / Projects (same core)
    # =========================
    def _extract_summary_enhanced(self, text: str) -> str:
        summary_section = self._extract_section(text, self.headers["summary"])
        if summary_section:
            cleaned = self._clean_summary_text(summary_section)
            if len(cleaned) > 20:
                return cleaned
        return self._extract_summary_from_beginning(text)

    def _clean_summary_text(self, text: str) -> str:
        text = re.sub(r'^(Summary|Objective|Profile)[:\s]*', '', text, flags=re.I)
        return re.sub(r'\s+', ' ', text).strip()

    def _extract_summary_from_beginning(self, text: str) -> str:
        lines = text.split('\n')
        content_lines = []
        skip = [self._re_email, self._re_phone_seed, self._re_linkedin, self._re_github]
        for line in lines[:20]:
            line = line.strip()
            if not line:
                continue
            if any(p.search(line) for p in skip):
                continue
            if len(line.split()) < 3:
                continue
            content_lines.append(line)
        summary = ' '.join(content_lines[:3])
        if len(summary) > 500:
            summary = summary[:500] + '...'
        return summary

    def _extract_certifications_enhanced(self, text: str) -> List[str]:
        section = self._extract_section(text, self.headers["certifications"])
        if not section:
            return []
        certs = []
        items = re.split(r'[,;|•▪▫◦‣⁃\n]', section)
        for item in items:
            item = item.strip(' -•▪▫◦‣⁃\t\n')
            if self._is_valid_certification(item):
                certs.append(item)
        return certs

    def _is_valid_certification(self, text: str) -> bool:
        if not text or not (3 <= len(text) <= 100):
            return False
        invalid = [r'^\d{4}$', r'^(jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)\b']
        return not any(re.match(p, text, re.I) for p in invalid)

    def _extract_projects_enhanced(self, text: str) -> List[Dict]:
        section = self._extract_section(text, self.headers["projects"])
        if not section:
            return []
        blocks = self._split_project_blocks(section)
        projects = []
        for b in blocks:
            if len(b.strip()) < 20:
                continue
            p = self._parse_project_block(b)
            if p and self._is_valid_project(p):
                projects.append(p)
        return projects

    def _split_project_blocks(self, section: str) -> List[str]:
        separators = [
            r'\n\s*\n',
            r'\n(?=Project\s*:)',  # "Project:" lines
            r'\n(?=[A-Z][^,\n]{5,50}(?:\s*[-–—]\s*))',
        ]
        blocks = [section]
        for sep in separators:
            new_blocks = []
            for b in blocks:
                parts = re.split(sep, b, flags=re.I)
                new_blocks.extend([p.strip() for p in parts if p.strip()])
            blocks = new_blocks
        return blocks

    def _parse_project_block(self, block: str) -> Dict:
        lines = [ln.strip() for ln in block.split('\n') if ln.strip()]
        result = {
            "name": "",
            "description": "",
            "technologies": [],
            "duration": "",
            "url": "",
            "details": block.strip(),
        }
        if lines:
            name = lines[0]
            name = re.sub(r'^Project\s*:?\s*', '', name, flags=re.I)
            result["name"] = name.strip()

        tech_patterns = [
            r'(?:Technologies?|Tech\s+Stack|Tools?|Built\s+with)[:\s]*([^.\n]{10,100})',
            r'(?:Using|Utilized|Implemented\s+with)[:\s]*([^.\n]{10,100})',
        ]
        for p in tech_patterns:
            m = re.search(p, block, re.I)
            if m:
                tech_text = m.group(1)
                technologies = [t.strip() for t in re.split(r'[,;|&/]', tech_text)]
                result["technologies"] = [t for t in technologies if t and len(t) > 1]
                break

        desc_lines = []
        for line in lines[1:]:
            if not any(k in line.lower() for k in ['technolog', 'tool', 'built']):
                desc_lines.append(line)
        result["description"] = ' '.join(desc_lines[:3])

        url_match = re.search(r'https?://[^\s,;]+', block)
        if url_match:
            result["url"] = url_match.group(0).rstrip(').,;:')

        return result

    def _is_valid_project(self, p: Dict) -> bool:
        if not p.get('name'):
            return False
        if not p.get('description') and not p.get('technologies'):
            return False
        return True

    # =========================
    # Languages (improved)
    # =========================
    def _extract_languages_enhanced(self, text: str) -> List[str]:
        section = self._extract_section(text, self.headers["languages"])
        if not section:
            return []
        return self._parse_languages_chunk(section)

    def _extract_languages_from_anywhere(self, text: str) -> List[str]:
        """
        Global fallback: find "Languages:" anywhere, or match known spoken languages in the whole doc.
        """
        # Try "Languages: ..." style chunks
        m = re.search(r'(?i)languages?\s*:\s*([^\n]+)', text)
        if m:
            langs = self._parse_languages_chunk(m.group(1))
            if langs:
                return langs

        # Match any known language terms in the whole document
        known = {l.lower(): l for l in self.skill_library.get_skills_by_category("spoken_languages")}
        found = set()
        tokens = re.split(r'[,\n;|/()\[\]-]+', text)
        for t in tokens:
            t = t.strip()
            base = re.sub(r'\s*\([^)]*\)\s*', '', t)  # remove parenthetical proficiency
            base = re.sub(r'\s*[-–—]\s*(fluent|native|basic|intermediate|advanced)\b.*', '', base, flags=re.I)
            if base.lower() in known:
                found.add(known[base.lower()])
        return sorted(found, key=lambda x: x.lower())

    def _parse_languages_chunk(self, chunk: str) -> List[str]:
        known = {l.lower(): l for l in self.skill_library.get_skills_by_category("spoken_languages")}
        items = re.split(r'[,;|•▪▫◦‣⁃\n]', chunk)
        out = []
        for item in items:
            item = item.strip(' -•▪▫◦‣⁃\t\n')
            item = re.sub(r'\s*\([^)]*\)\s*', '', item)
            item = re.sub(r'\s*[-–—]\s*(fluent|native|basic|intermediate|advanced)\b.*', '', item, flags=re.I)
            if not item:
                continue
            if item.lower() in known:
                out.append(known[item.lower()])
            elif self._looks_like_language(item):
                out.append(item.title())
        return sorted(set(out), key=lambda x: x.lower())

    def _looks_like_language(self, text: str) -> bool:
        if not text or not (3 <= len(text) <= 20):
            return False
        if not text.replace(' ', '').isalpha():
            return False
        return True

    # =========================
    # Section Detection / Utilities
    # =========================
    def _extract_section(self, text: str, headers: List[str]) -> str:
        if not text or not headers:
            return ""
        variants = []
        for h in headers:
            variants.extend([h, h.upper(), h.title(), h.replace(' ', ''), h.replace(' ', '_')])
        pattern = r'(?mi)^\s*(?:' + '|'.join(re.escape(h) for h in variants) + r')\s*[:\-–]?\s*'
        start = re.search(pattern, text)
        if not start:
            return ""
        start_pos = start.end()
        remaining = text[start_pos:]
        # Next section: all-caps or typical header
        next_section_pattern = r'(?mi)^\s*[A-Z][A-Z\s&/]{5,40}\s*[:\-–]?\s*$'
        nxt = re.search(next_section_pattern, remaining)
        if nxt:
            end_pos = start_pos + nxt.start()
            section_text = text[start_pos:end_pos]
        else:
            section_text = remaining
        return section_text.strip()

    def _clean_text(self, text: str) -> str:
        if not text:
            return ""
        text = text.replace('\u2022', '•')  # bullets
        text = text.replace('\u2013', '-')  # en dash
        text = text.replace('\u2014', '—')  # em dash
        text = text.replace('\u2019', "'")
        text = text.replace('\u201c', '"')
        text = text.replace('\u201d', '"')
        text = self._re_cr.sub('\n', text)
        text = self._re_newlines.sub('\n\n', text)
        text = self._re_spaces.sub(' ', text)
        return text.strip()

    @lru_cache(maxsize=1000)
    def _cached_nlp(self, text_hash: str, text: str):
        return self.nlp(text)

    def _get_text_hash(self, text: str) -> str:
        return hashlib.md5(text.encode()).hexdigest()[:16]

    # =========================
    # Batch Utilities
    # =========================
    def clear_cache(self):
        self._cache.clear()

    def get_processing_stats(self) -> Dict:
        return {
            "cache_size": len(self._cache),
            "max_cache_size": self.CACHE_SIZE,
            "processing_mode": self.processing_mode,
            "max_workers": self.MAX_WORKERS,
            "batch_size": self.BATCH_SIZE,
        }

    async def process_single_resume(self, file_path: str) -> Dict:
        return await self._process_single_async(file_path)

    def process_single_resume_sync(self, file_path: str) -> Dict:
        return self._process_single_sync(file_path)


# ---------------------------
# Factory Functions
# ---------------------------
def create_resume_parser(processing_mode: str = EnhancedResumeParser.MULTI_THREADED) -> EnhancedResumeParser:
    return EnhancedResumeParser(processing_mode=processing_mode)

def create_batch_parser(max_workers: int = None) -> EnhancedResumeParser:
    parser = EnhancedResumeParser(processing_mode=EnhancedResumeParser.ASYNC_BATCH)
    if max_workers:
        parser.MAX_WORKERS = min(max_workers, cpu_count() or 2)
    return parser

# Maintain backward compatibility
ResumeParser = EnhancedResumeParser
